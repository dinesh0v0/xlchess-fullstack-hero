from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_doc():
    doc = Document()
    
    title = doc.add_heading('DAChess: Comprehensive QA & Testing Strategy', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph("Version: 1.0.0 | Orand Gaming Final Review").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Functional Testing
    doc.add_heading('1. Functional Testing (End-to-End)', level=1)
    doc.add_paragraph("Objective: Validate that user interactions trigger the correct game logic seamlessly.")
    doc.add_paragraph(
        "Solution: We have integrated @playwright/test to perform full-flow End-to-End (E2E) testing. "
        "Playwright automates browser actions (clicking pieces, dragging, toggling menus) to simulate a real human "
        "playing the game. Our playwright.config.ts is established to run these functional validations on every deployment."
    )
    
    # 2. Cross-Browser & Responsive
    doc.add_heading('2. Cross-Browser & Responsive Testing', level=1)
    doc.add_paragraph("Objective: Guarantee UI integrity across various engines and viewports.")
    doc.add_paragraph(
        "Solution: Tailwind CSS ensures structural responsiveness. For automated verification, the Playwright testing matrix "
        "has been configured to run tests simultaneously across Chromium (Chrome/Edge), WebKit (Safari), and Firefox. "
        "Additionally, it tests mobile viewports (e.g., iPhone 13) to ensure the chessboard resizes without clipping."
    )
    
    # 3. Performance Testing
    doc.add_heading('3. Performance Testing', level=1)
    doc.add_paragraph("Objective: Achieve 60fps animations and rapid TTFB (Time to First Byte).")
    doc.add_paragraph(
        "Solution: Vite minimizes bundle size while Framer Motion handles GPU-accelerated layout transitions. "
        "To monitor this formally, we run Google Lighthouse CI during the build process targeting the following thresholds:\n"
        "- LCP (Largest Contentful Paint): < 2.5s\n"
        "- CLS (Cumulative Layout Shift): < 0.1\n"
        "- FID (First Input Delay): < 100ms"
    )
    
    # 4. Security Testing
    doc.add_heading('4. Security Testing', level=1)
    doc.add_paragraph("Objective: Prevent vulnerabilities like XSS, Dependency Exploits, and Injection.")
    doc.add_paragraph(
        "Solution: We actively run `npm audit` in the CI/CD pipeline. The latest audit returned 0 vulnerabilities. "
        "Once the Django REST backend is deployed, strict Content Security Policy (CSP) headers will be enforced to "
        "neutralize Cross-Site Scripting, and Django's built-in CSRF tokens will secure all API requests."
    )
    
    # 5. UAT & Bug Tracking
    doc.add_heading('5. User Acceptance Testing (UAT) & Bug Tracking', level=1)
    doc.add_paragraph("Objective: Solicit beta-tester feedback systematically.")
    doc.add_paragraph(
        "Solution: Beta testers will perform exploratory testing. All bugs will be tracked via standard GitHub Issues "
        "using the following schema:"
    )
    doc.add_paragraph("Bug Tracking Schema Template:", style='List Bullet')
    doc.add_paragraph("Title: [Component] Brief description", style='List Bullet 2')
    doc.add_paragraph("Steps to Reproduce: 1, 2, 3...", style='List Bullet 2')
    doc.add_paragraph("Expected vs Actual Behavior", style='List Bullet 2')
    doc.add_paragraph("Environment: (Browser, OS, Viewport)", style='List Bullet 2')
    
    doc.save('QA_Testing_Solutions.docx')
    print("Successfully created QA_Testing_Solutions.docx")

if __name__ == '__main__':
    create_doc()
